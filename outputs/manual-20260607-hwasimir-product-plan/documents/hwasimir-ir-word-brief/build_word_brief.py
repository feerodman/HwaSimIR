from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(r"D:\HwaSimIR\outputs\manual-20260607-hwasimir-product-plan\documents\hwasimir-ir-word-brief")
OUT_DOCX = OUT_DIR / "HwaSimIR_IR_Product_Roadmap_Text.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(95, 107, 122)
FILL = "F4F6F9"
LIGHT_BLUE = "E8F0FB"
GREEN_FILL = "EAF7F2"
ORANGE_FILL = "FFF3E5"


def set_run_font(run, name="Calibri", east="微软雅黑", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_font(paragraph, size=11, color=INK, bold=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_cell_text(cell, text, bold=False, color=INK, size=9.5, fill=None, align=None):
    if fill:
        set_cell_shading(cell, fill)
    cell.text = ""
    for i, part in enumerate(str(text).split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        if align:
            p.alignment = align
        run = p.add_run(part)
        set_run_font(run, size=size, color=color, bold=bold)


def add_paragraph(doc, text, size=11, color=INK, bold=False, after=8, before=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.333
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
    else:
        set_run_font(run, size=12, color=DARK_BLUE, bold=True)
    return p


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11.5, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.25
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_matrix(doc, headers, rows, widths, header_fill=FILL, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        set_cell_text(table.cell(0, i), header, bold=True, color=DARK_BLUE, size=9.6, fill=header_fill, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size, fill=None, align=WD_ALIGN_PARAGRAPH.LEFT if i else WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        r = p.add_run(item)
        set_run_font(r, size=10.5, color=INK)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("HwaSimIR 红外成像仿真产品规划文字稿")
    set_run_font(r, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("北京华力创通科技股份有限公司 | 内部讨论稿")
    set_run_font(r, size=9, color=MUTED)
    return doc


def build():
    doc = setup_document()

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("国产化三维红外成像仿真产品规划文字稿")
    set_run_font(r, size=24, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("基于国产化 ARM 板卡、开源 3D 引擎与集团仿真资源的产品化平台及行业解决方案")
    set_run_font(r, size=12.5, color=MUTED)

    add_callout(
        doc,
        "建议立项口径",
        "本方向不宜定位为一次性项目交付，也不宜只做单点红外渲染 demo。更合适的定位是“国产化三维红外成像仿真产品平台”，以军工、低空经济、商业航天解决方案包作为首批商业化抓手，后续逐步接入 AFSIM、VR、虚幻 5、雷达成像和 AI 数据生成能力，形成公司传感器仿真产品线的一部分。",
    )

    meta = doc.add_table(rows=4, cols=2)
    set_table_geometry(meta, [2200, 7160])
    meta_rows = [
        ("文档定位", "内部方向讨论稿 / 产品规划文字稿"),
        ("对应 PPT", "HwaSimIR_IR_Product_Roadmap.pptx"),
        ("主要来源", "申报表 V1.0、近期 HwaSimIR 工程积累、产品规划讨论要求"),
        ("日期", "2026 年 06 月 07 日"),
    ]
    for i, (k, v) in enumerate(meta_rows):
        set_cell_text(meta.cell(i, 0), k, bold=True, color=DARK_BLUE, size=9.5, fill=FILL, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(meta.cell(i, 1), v, size=9.5)

    doc.add_page_break()

    add_heading(doc, "一、总体判断：产品平台为主，行业解决方案为抓手", 1)
    add_paragraph(
        doc,
        "三维场景红外成像仿真产品的核心价值，不只是把三维场景渲染成红外图像，而是为红外设备测试、目标识别算法验证、半实物仿真、训练评估和任务推演提供一个可重复、可配置、可解释的数字试验环境。若只把它理解为一个项目中的功能模块，后续容易陷入按客户需求反复定制、难以复用的状态；若只把它理解为一个通用软件产品，又容易脱离军工、低空、航天客户实际采购方式。因此建议采用“平台产品 + 行业方案包”的双层定位。",
    )
    add_paragraph(
        doc,
        "平台产品沉淀的是共性能力，包括红外物理模型、三维场景引擎、国产化 ARM 运行时、SDK/API、GUI 配置、模型库、材料库、测试报告和自动化验证工具链。行业方案包则面向具体客户场景进行包装，例如军工半实物测试方案、低空无人机红外感知验证方案、商业航天载荷仿真方案、安防应急演训方案。这样既能形成可销售的产品形态，也能保留解决方案交付所需的灵活性。",
    )
    add_callout(
        doc,
        "一句话定位",
        "面向军工、低空经济和商业航天的国产化三维红外成像仿真平台，提供从红外物理建模、场景生成、传感器成像到算法评估和系统集成的完整工具链。",
        fill=GREEN_FILL,
    )

    add_heading(doc, "二、产品介绍：产品是什么，解决什么问题", 1)
    add_heading(doc, "2.1 产品基本形态", 2)
    add_paragraph(
        doc,
        "本产品面向三维场景中的红外成像仿真。用户输入三维场景、目标模型、材料参数、温度状态、大气条件、传感器参数和任务脚本，系统输出不同波段下的红外图像、图像序列、目标状态日志、仿真参数记录和测试报告。产品既可以作为独立工具使用，也可以作为仿真系统中的红外传感器插件接入 AFSIM、VR、虚幻 5 或半实物平台。",
    )
    add_paragraph(
        doc,
        "与传统视景软件相比，本产品的重点不是通用可见光画面，而是红外成像链路的可解释性和工程可交付性。红外图像的生成需要考虑目标温度、材料发射率、反射率、大气透过率、传感器波段、热源和亮点等因素。如果这些因素不能被配置、记录和回放，就很难支撑军工客户的试验验证和算法评估。因此，产品从一开始就应把“可配置、可追溯、可验收”作为核心设计原则。",
    )

    add_heading(doc, "2.2 产品与解决方案的关系", 2)
    rows = [
        ("平台产品", "红外仿真内核、GUI、SDK/API、国产 ARM 运行时、材料/目标/大气模型库、自动化测试与报告。", "形成可复用版本和产品授权，降低重复项目开发成本。"),
        ("行业方案包", "军工半实物测试、低空无人机红外感知、商业航天载荷验证、安防应急演训、科研教学实验。", "贴近客户预算和验收场景，带动平台能力持续迭代。"),
        ("长期平台化方向", "红外 + 雷达 + 可见光多模态仿真，接入 AFSIM、VR、虚幻 5、AI 数据生成与自动评估。", "形成集团级传感器仿真底座，支撑低空经济和商业航天新赛道。"),
    ]
    add_matrix(
        doc,
        ["形态", "主要内容", "商业意义"],
        rows,
        [1700, 4400, 3260],
        header_fill=LIGHT_BLUE,
        font_size=9.3,
    )

    add_heading(doc, "三、市场前景与具体用户", 1)
    add_paragraph(
        doc,
        "红外成像仿真的市场需求来自三个共同痛点：第一，实地试验成本高且组织周期长，尤其是军工装备、靶场试验、无人机低空场景和航天载荷验证；第二，真实场景不可完全复现，天气、时间、目标状态和背景环境都会影响测试结果；第三，国产化替代和自主可控要求提高，客户不希望核心仿真链路长期依赖国外闭源软件或单一 X86 环境。上述痛点决定了该产品不只是研发工具，也可以成为测试、训练、评估和数据生成的基础设施。",
    )
    rows = [
        ("军工装备", "军工研究所、总体单位、红外设备厂商、靶场/试验基地", "红外导引、目标识别、半实物仿真、实验室测试、靶场前置验证。", "第一优先级。需求刚性强、预算相对明确、国产化约束强，适合定义 V1.0 验收标准。"),
        ("低空经济", "低空监管平台、无人机运营商、机场/园区/港口、反无人机集成商", "无人机目标识别、红外/可见光/雷达融合感知、低空安全数字孪生。", "第二优先级。市场增长快，但需求分散，适合打包场景库和解决方案。"),
        ("商业航天", "卫星总体单位、载荷研制单位、商业航天公司、地面仿真团队", "载荷算法上星前验证、星地红外场景、空间目标和地物红外背景仿真。", "战略储备方向。可与公司航天业务、任务仿真和载荷验证能力协同。"),
        ("安防应急与工业", "边海防、能源电力、消防应急、重点设施安防、工业检测企业", "夜间、烟雾、复杂天气、热异常目标和监控算法评测。", "适合作为产品成熟后的横向复制市场。"),
        ("科研教育与内部应用", "高校、科研院所、公司内部测试部门", "教学实验、算法研究、模型验证、内部降本增效。", "适合作为早期反馈和模型库建设场景。"),
    ]
    add_matrix(
        doc,
        ["市场方向", "具体用户", "典型需求", "切入判断"],
        rows,
        [1300, 2100, 3100, 2860],
        header_fill=FILL,
        font_size=8.7,
    )
    add_heading(doc, "3.1 三类首批场景的讲述方式", 2)
    add_paragraph(
        doc,
        "军工场景可以这样讲：客户真正需要的不是一张好看的红外图，而是在实验室里尽量复现靶场前置试验条件。红外设备厂商、导引头算法团队或总体单位可以通过本产品配置目标、背景、距离、姿态、发动机状态、命中状态和大气条件，生成可重复的红外图像序列，用于算法验证、设备联调和试验方案预演。这样做的价值在于把原本必须到靶场才能发现的问题前置到实验室阶段，降低试验成本和迭代周期。",
    )
    add_paragraph(
        doc,
        "低空经济场景可以这样讲：无人机低空监管和反制系统需要在夜间、逆光、烟雾、复杂背景和不同天气条件下识别小目标，真实数据采集成本高、覆盖不全且存在安全风险。红外仿真平台可以生成不同高度、速度、姿态、载荷和背景下的无人机红外数据，并与雷达、可见光、电子围栏或管控平台联动，服务低空安全方案验证和算法训练。",
    )
    add_paragraph(
        doc,
        "商业航天场景可以这样讲：卫星载荷和星地仿真团队在上星前需要评估传感器在不同轨道、观测角度、地物背景和大气条件下的成像效果。红外仿真产品可以作为载荷算法和地面仿真系统之间的成像通道，提供星地红外场景、目标热特征和图像序列，帮助商业航天客户降低上星前验证风险。该方向短期可能不如军工试点直接，但与公司商业航天战略高度相关，适合作为第二阶段重点拓展。",
    )

    add_heading(doc, "四、核心技术点与已有基础", 1)
    add_paragraph(
        doc,
        "HwaSimIR 当前已经具备产品化起步的技术基础。仓库中包含 C++、Qt、Panda3D、OpenCV、Eigen 等工程基础，能够支撑桌面工具、三维场景加载、图像处理和原生性能优化。近期工作进一步把红外成像链路拆解成可验证的阶段：Stage 3 保持 MODTRAN tau-only 的受控实验链路，用于大气透过率加载与调试；Stage 4 将发动机热斑 ThermalHotspot 与武器命中亮点 BrightSpot 分离，并通过 targetType + targetPlatID + targetID 的全协议键解决目标映射问题。这些工作虽然还不是完整产品，但已经体现出一个重要方向：产品不能只看渲染结果，还要能解释每个红外输出来自哪些输入状态。",
    )
    add_heading(doc, "4.1 红外物理模型", 2)
    add_paragraph(
        doc,
        "红外物理模型是产品的核心。它需要覆盖目标温度场、材料发射率、反射率、大气透过率、传感器波段、背景辐射和局部热源等因素。V1.0 阶段建议先形成工程可用的最小闭环：支持短波、中波、长波三个典型波段，支持目标基础温度与局部热斑配置，支持大气透过率对成像灰度的影响，并将关键输入和输出记录到日志中。后续再逐步接入更精细的辐亮度链路、太阳反射、天空背景、路径辐射和传感器噪声模型。",
    )
    add_heading(doc, "4.2 三维场景与渲染引擎", 2)
    add_paragraph(
        doc,
        "Panda3D 适合作为早期轻量化和国产化 ARM 适配路线，因为它开源、可控、工程成本较低，适合把红外模型嵌入到现有 C++/Python 生态中。与此同时，虚幻 5 不应被排除在产品路线之外。UE5 更适合作为下一代高保真场景、大规模地形、复杂光照和沉浸式交互路线，尤其适用于低空城市/园区场景、训练演示、客户展示和高质量数据生成。建议形成双路线：Panda3D/ARM 负责轻量部署和工程交付，UE5 负责高保真场景和长期视觉竞争力。",
    )
    add_heading(doc, "4.3 目标状态、协议与可验证性", 2)
    add_paragraph(
        doc,
        "红外仿真产品要接入任务系统，就必须处理目标状态、平台编号、武器状态、视线状态和传感器参数。近期 HwaSimIR 工作中，目标映射已经从单一 targetID 扩展到 targetType + targetPlatID + targetID 的完整协议键，这一点非常重要。它意味着系统可以区分不同平台、不同类型下的同名目标，避免错误目标被看向、命中或触发亮点。产品化时应继续沿着这个方向，把协议适配层作为独立模块，支持 UDP、二进制协议、JSON、DIS/HLA 或 AFSIM 接口。",
    )
    add_callout(
        doc,
        "已有基础转产品的关键",
        "从工程角度看，现阶段不是简单补 UI，而是要把红外模型、目标协议、运行配置、日志诊断、自动化 smoke 测试和验收指标打包成产品能力。只有做到“输入可配置、过程可追踪、结果可复现、问题可定位”，才能支撑军工客户验收。",
    )

    add_heading(doc, "五、产品模块设计", 1)
    rows = [
        ("场景与资产模块", "管理地形、建筑、目标模型、材质、纹理和场景脚本。支持导入公司已有三维模型素材，并逐步形成低空、航天、靶场、安防等标准场景库。", "场景加载、LOD、坐标系统、模型格式转换、资产版本管理。", "场景文件、目标实例、背景环境和批量试验工况。"),
        ("红外物理模块", "负责温度、材料、大气、波段、热斑、亮点和最终红外灰度/辐亮度计算。该模块决定产品是否具备物理可信度。", "温度场模型、发射率/反射率、MODTRAN tau、热点/亮点分离、辐亮度链路。", "红外图像、图像序列、分项灰度/辐亮度日志。"),
        ("传感器模块", "描述红外相机或载荷的波段、视场角、分辨率、帧率、噪声、动态范围和成像模式。", "相机模型、视线控制、内外参、传感器噪声、图像输出管线。", "传感器配置、红外图像流、算法测试输入。"),
        ("协议与系统集成模块", "负责接入外部仿真系统、半实物平台和任务推演系统，使红外仿真能够随目标状态实时变化。", "UDP/TCP、二进制协议、JSON、DIS/HLA、AFSIM 适配、状态同步。", "目标状态、武器状态、视线状态、事件触发记录。"),
        ("GUI 与试验管理模块", "提供可视化配置、场景编辑、参数管理、批量试验、结果查看和报告导出，降低售前和交付使用门槛。", "Qt/桌面 GUI、参数模板、批处理、日志解析、报告生成。", "可操作软件界面、试验任务、结果报告。"),
        ("跨平台运行模块", "支持 Windows 开发调试、Linux 部署、国产 ARM 板卡运行，并预留 GPU/FPGA 加速接口。", "CMake/Qt 工程化、Panda3D/UE5 运行时、ARM GPU、性能监控。", "部署包、运行时、性能报告和交付文档。"),
        ("AI 与数据模块", "在产品成熟后提供场景生成、参数反演、红外合成数据集、自动测试用例和报告生成能力。", "生成式 AI、仿真参数推荐、数据标注、模型评估、异常诊断。", "训练数据集、评测报告、自动化工况集。"),
    ]
    add_matrix(
        doc,
        ["模块", "职责说明", "关键技术", "输出物"],
        rows,
        [1450, 3500, 2550, 1860],
        header_fill=LIGHT_BLUE,
        font_size=8.2,
    )
    add_heading(doc, "5.1 典型使用流程", 2)
    add_paragraph(
        doc,
        "从用户视角看，产品的使用流程可以分为六步。第一步是建立场景，用户选择靶场、机场、园区、城市低空、轨道观测或海面等基础环境，导入地形、建筑、目标模型和材质。第二步是配置红外参数，设置目标温度、材料发射率、局部热斑、大气条件、波段和传感器参数。第三步是配置任务状态，例如目标运动轨迹、平台姿态、发动机状态、武器事件、视线角和采样频率。第四步是运行仿真，系统输出红外图像、视频流或图像序列，并同步输出状态日志。第五步是结果评估，用户可以对比算法识别结果、灰度变化、目标可见性和异常状态。第六步是导出报告，用于项目验收、算法评估或试验方案评审。",
    )
    add_heading(doc, "5.2 模块展开说明", 2)
    add_paragraph(
        doc,
        "场景与资产模块是产品可复用程度的基础。早期可以从公司已有模型素材和典型靶场/飞行器/低空场景开始，逐步形成标准化资产包。资产库不只是模型文件集合，还应记录模型尺度、坐标基准、材质参数、红外属性和版本信息，否则后续很难复现实验结果。",
    )
    add_paragraph(
        doc,
        "红外物理模块是产品可信度的核心。该模块需要把目标主体、局部热源、背景、大气和传感器输出拆开建模，并通过日志保存每一类贡献。近期 HwaSimIR 中热点和亮点分离的工作非常有价值，因为它避免把发动机状态、武器命中和泛化毁伤状态混在一起。产品化时也应保持这种边界：发动机状态控制尾部热斑，命中/打击事件控制亮点，材料和温度模型控制主体辐射。",
    )
    add_paragraph(
        doc,
        "传感器模块决定产品能否接近真实设备。不同红外设备的波段、分辨率、视场角、曝光策略、动态范围和噪声特性不同，如果系统只能输出一类固定灰度图，就很难服务不同客户。建议 V1.0 先支持常用波段、视场角、分辨率和图像输出参数，后续增加传感器噪声、非均匀性校正、动态范围压缩和成像链路延迟。",
    )
    add_paragraph(
        doc,
        "协议与系统集成模块是把产品从独立工具变成平台能力的关键。它需要把外部系统输入的目标状态、平台编号、武器事件和相机视线转换为内部仿真状态，也需要把红外图像、目标可见性和仿真日志输出给算法系统或半实物设备。近期目标全协议键的改造说明，协议层不能简单依赖 targetID，而要考虑目标类型、平台编号、目标编号等组合键。",
    )
    add_paragraph(
        doc,
        "GUI 与试验管理模块决定产品是否能被售前、测试和客户工程人员使用。GUI 不应只做参数输入框，而应支持场景模板、参数模板、批量工况、结果浏览、日志筛选和报告导出。对于军工客户而言，报告中应明确列出场景、目标、传感器、大气、波段、关键状态和输出结果，这些内容可以直接用于验收材料。",
    )
    add_paragraph(
        doc,
        "跨平台运行模块决定产品能否符合国产化和嵌入式部署要求。Windows 环境适合研发和售前演示，Linux 适合服务器和工程交付，国产 ARM 板卡适合便携式终端和信创环境。产品路线应明确哪些能力必须在 ARM 端实时运行，哪些能力可以离线生成，哪些能力可以由服务器或 GPU 工作站承担。",
    )
    add_paragraph(
        doc,
        "AI 与数据模块可以放在第二阶段以后重点建设。早期不要为了追 AI 概念而影响 V1.0 闭环，但要在数据结构和日志体系上预留接口。只要场景参数、目标状态和输出图像能够标准化保存，后续就可以自然扩展到场景生成、参数反演、合成数据集和自动评测。",
    )

    add_heading(doc, "六、与 AFSIM、VR、虚幻 5、雷达成像等软件的集成路线", 1)
    add_paragraph(
        doc,
        "该产品不应作为孤立红外工具存在。公司已有 AFSIM、VR、虚幻 5、雷达成像等方向，恰好可以形成任务级仿真、沉浸式训练、高保真视景和多模态传感器仿真的组合能力。红外仿真产品在其中承担“红外传感器通道”和“红外数据生成通道”的角色。集成路线应从接口松耦合开始，再逐步走向场景、时间、坐标和事件的一体化同步。",
    )
    rows = [
        ("AFSIM", "AFSIM 负责作战任务、平台行为、航迹、传感器调度和战术事件；HwaSimIR 提供红外传感器成像结果。", "先通过状态文件/UDP 接入目标位置、姿态、平台编号和事件，再扩展到 DIS/HLA 或插件式接口。", "形成“任务推演 + 红外成像”的联合仿真能力，用于导引、识别、态势评估和方案论证。"),
        ("VR", "VR 负责沉浸式操作、训练展示和复盘交互；红外仿真提供传感器视角、夜间/烟雾/热目标等训练画面。", "通过共享场景、相机视角、事件时间轴和回放数据实现联动。", "形成训练演示、装备操作、低空安全处置和复盘评估方案。"),
        ("虚幻 5", "UE5 负责高保真场景、大规模地形、复杂城市/园区环境和客户展示效果。", "建议保留 Panda3D/ARM 轻量路线，同时建立 UE5 高保真分支，红外模型以材质、后处理或插件形式迁移。", "提升视觉表现和大场景能力，适合低空经济、商业航天展示和高质量数据生成。"),
        ("雷达成像", "雷达仿真负责距离、多普勒、散射、SAR/ISAR 或目标探测链路；红外仿真负责热特征和图像感知。", "统一目标库、场景坐标、时间戳和传感器配置，输出多模态数据包。", "形成红外 + 雷达 + 可见光的多模态感知评估平台，服务低空监管和目标识别算法。"),
        ("半实物/硬件在环", "外部硬件设备、导引头、图像处理板卡或算法服务器需要实时图像流和状态闭环。", "提供图像序列、实时视频流、同步触发、延迟监控和状态日志。", "支撑实验室测试、靶场前置验证和设备验收。"),
    ]
    add_matrix(
        doc,
        ["集成对象", "分工关系", "集成方式", "产品价值"],
        rows,
        [1200, 3000, 2800, 2360],
        header_fill=FILL,
        font_size=8.35,
    )
    add_heading(doc, "6.1 集成节奏建议", 2)
    add_paragraph(
        doc,
        "集成工作应避免一开始就追求大而全。第一阶段建议采用松耦合方式：通过文件、UDP 或简单 API 传递目标状态和传感器参数，让红外仿真先能被外部系统调用。第二阶段再统一坐标系、时间戳、事件触发和场景资产，实现与 AFSIM、VR 或 UE5 的同步联动。第三阶段才考虑更深的插件化、DIS/HLA、实时图像流和硬件在环接口。这样做可以降低早期集成风险，也能让产品先形成独立价值。",
    )
    add_paragraph(
        doc,
        "与虚幻 5 的关系尤其需要讲清楚。虚幻 5 不是替代现有 Panda3D/ARM 路线，而是面向高保真、大场景、沉浸式展示和数据生成的增强路线。早期产品仍应保证轻量可部署，尤其是国产 ARM 和离线环境；当需要做低空城市级场景、客户演示或高质量数据集时，再把 UE5 作为高保真场景引擎接入。这样既保住产品交付节奏，也给未来视觉表现留下上限。",
    )
    add_paragraph(
        doc,
        "与雷达成像的关系也应定位为互补。红外擅长热特征和图像识别，雷达擅长距离、速度、全天候探测和散射特征。低空监管、无人机反制和复杂战场环境往往需要多传感器融合，因此红外仿真与雷达成像一旦共享同一套目标库、场景坐标和时间线，就可以形成更高价值的融合评估平台。",
    )

    add_heading(doc, "七、创新点、难点与行业壁垒", 1)
    add_paragraph(
        doc,
        "本方向的创新点不应只表述为“国产 ARM + 开源 3D 引擎”，还应进一步强调它在红外物理链路、国产化部署、可解释验证和集团级仿真集成上的组合创新。单一技术点容易被模仿，但把红外模型、国产运行环境、协议接入、模型库、数据校准和自动化验收体系串起来，就会形成更难复制的工程壁垒。",
    )
    add_heading(doc, "7.1 核心创新点", 2)
    add_bullets(
        doc,
        [
            "国产化 ARM 板卡红外仿真适配：面向国产硬件、便携终端和信创环境，降低对国外软件和 X86 工作站的依赖。",
            "开源 3D 引擎红外链路集成：在 Panda3D 等开源引擎上集成红外物理模型，具备源码可控、成本可控和可定制优势。",
            "多波段红外与局部状态建模：支持短波、中波、长波方向，并能区分发动机热斑、武器命中亮点、目标主体和背景。",
            "协议化目标状态联动：通过完整目标键、武器状态、视线状态和日志诊断，为 AFSIM、半实物和任务系统接入打基础。",
            "自动化验证与可追溯交付：通过严格检查、smoke 测试、日志和报告，让客户不仅看到画面，还能看到结果依据。",
        ],
    )
    add_heading(doc, "7.2 主要难点", 2)
    add_paragraph(
        doc,
        "第一类难点是性能。红外计算如果逐像素叠加温度、材质、大气、反射和局部热源，在 ARM 端容易影响实时帧率，需要通过 LOD、GPU 并行、分层渲染、缓存和分辨率策略进行优化。第二类难点是精度。红外仿真不能只凭视觉主观判断，需要实测数据、材料参数和传感器参数校准，否则难以通过军工客户验收。第三类难点是工程化。客户环境复杂，可能涉及 Windows、Linux、国产 ARM、离线部署、保密网络和硬件在环设备，产品必须具备稳定部署和问题定位能力。",
    )
    add_heading(doc, "7.3 壁垒建设方式", 2)
    add_paragraph(
        doc,
        "建议把壁垒建设分为四类：第一，模型壁垒，持续积累材料、目标、温度、大气和传感器模型库；第二，数据壁垒，联合试验单位积累实测校准数据和仿真-实测对照样本；第三，工程壁垒，形成 GUI、SDK、部署包、日志体系、自动化测试和验收规范；第四，生态壁垒，把红外仿真接入 AFSIM、VR、UE5、雷达成像等公司已有能力，形成组合方案。",
    )

    add_heading(doc, "八、对标公司与差异化策略", 1)
    add_paragraph(
        doc,
        "对标时不宜只找完全同类产品，因为红外成像仿真横跨视景仿真、热红外签名、传感器仿真、任务推演和航天任务分析等多个领域。建议按能力拆分对标：视景平台对标 Presagis/CAE，热红外物理对标 ThermoAnalytics，EO/IR 传感器仿真对标 OKTAL-SE，联合仿真接口对标 MAK，航天任务与载荷分析对标 Ansys/AGI。我们的差异化不在于一次性覆盖所有功能，而在于国产化部署、低成本定制、公司内部仿真体系集成和面向低空/航天/军工场景的本地化交付。",
    )
    rows = [
        ("Presagis / CAE", "Vega Prime / Creator 等视景仿真工具链", "成熟场景生态、视景工具链、仿真工程经验。", "以国产 ARM、开源可控、红外专用链路切入，降低授权和禁运风险。"),
        ("ThermoAnalytics", "热模型与红外签名仿真", "热传导、材料参数、红外签名工程化能力。", "结合国产运行环境和半实物测试流程，做更贴近国内交付的轻量方案。"),
        ("OKTAL-SE", "EO/IR 传感器场景仿真", "传感器级建模、任务场景、工程验证流程。", "围绕低空、航天、军工建立本地化场景库和可交付方案。"),
        ("MAK Technologies", "VR-Forces / VR-Vantage 联合仿真", "DIS/HLA、任务级仿真和系统集成。", "接入公司 AFSIM/VR 资源，补齐红外感知通道。"),
        ("Ansys / AGI", "STK 航天任务分析与载荷链路", "轨道、星座、载荷和链路分析。", "在商业航天场景中补充国产红外成像仿真和数据生成能力。"),
    ]
    add_matrix(doc, ["对标对象", "代表方向", "可学习点", "差异化机会"], rows, [1500, 2200, 2700, 2960], font_size=8.35)

    add_heading(doc, "九、实现路径与阶段规划", 1)
    add_paragraph(
        doc,
        "实现路径建议分三阶段推进，不把所有能力压到 V1.0。第一阶段先把红外仿真内核做成可交付产品，重点解决 GUI、接口、ARM 适配、模型库雏形和试点客户验收；第二阶段扩展行业方案包，重点围绕低空经济、军工半实物和商业航天建立可复制场景；第三阶段形成多模态传感器仿真平台，接入雷达、可见光、AI 数据生成和集团级任务仿真体系。",
    )
    rows = [
        ("阶段一：2026-2027", "V1.0/V1.5", "红外内核产品化、GUI 原型、SDK/API、国产 ARM 适配、军工试点客户。", "形成可演示、可部署、可验收的最小产品闭环。"),
        ("阶段二：2028", "V2.0", "低空无人机场景包、商业航天载荷仿真包、AFSIM/VR 接口、UE5 高保真分支、模型库商品化。", "形成可复制行业解决方案和标杆案例。"),
        ("阶段三：2029-2030", "V3.0", "红外 + 雷达 + 可见光多模态仿真，AI 合成数据服务，跨平台规模化交付。", "形成集团级传感器仿真平台和持续收入产品线。"),
    ]
    add_matrix(doc, ["阶段", "版本", "建设重点", "阶段成果"], rows, [1700, 1200, 4300, 2160], header_fill=LIGHT_BLUE)

    add_heading(doc, "十、投入产出与资源需求", 1)
    add_paragraph(
        doc,
        "投入产出需要按“核心团队 + 项目牵引 + 产品沉淀”的逻辑理解。首年不建议搭建过大的固定团队，而应以 5-6 人核心团队形成 V1.0 最小闭环；第二、三年结合试点客户和行业方案扩展到 8-10 人；第四、五年如果产品复制和多模态平台路线成立，再扩展到 10-12 人左右。固定资产投入主要集中在国产 ARM 板卡、红外相机、测试工位、数据存储、模型库和校准环境。",
    )
    rows = [
        ("第 1 年", "5-6 人", "ARM 板卡、红外相机、测试工位、基础模型库", "内部节支、原型验证、2-3 家试点客户需求定义。", "形成 V1.0、软著申请、试点演示材料。"),
        ("第 2-3 年", "8-10 人", "测试环境扩展、数据/模型库、客户交付环境", "试点销售、军工半实物方案、低空/航天方案复制。", "形成 3 年可销售版本、标杆案例、专利/技术秘密。"),
        ("第 4-5 年", "10-12 人", "多模态测试环境、AI 数据生成平台、规模化交付工具", "产品授权、行业解决方案、维护服务、模型库服务、AI 数据服务。", "形成传感器仿真产品线和集团平台复用能力。"),
    ]
    add_matrix(doc, ["阶段", "人员投入", "固定资产/数据投入", "产出方式", "沉淀资产"], rows, [1100, 1300, 2500, 2500, 1960], header_fill=FILL, font_size=8.45)
    add_callout(
        doc,
        "经营逻辑",
        "前期以军工和低空项目验证付费能力，中期把 GUI、接口、模型库和部署包沉淀成可复制产品，后期通过产品授权、行业解决方案、年度维护、模型库服务和 AI 数据生成服务形成复合收入。",
        fill=ORANGE_FILL,
    )

    add_heading(doc, "十一、AI 化方向", 1)
    add_paragraph(
        doc,
        "AI 化不应停留在宣传口径，而应落在四个具体能力上。第一是场景自动生成，根据任务描述、区域类型、目标类型和天气条件生成低空、航天、安防或靶场场景，减少人工建模工作量。第二是参数反演与校准，利用实测红外图像反推材料、温度、大气和传感器参数，提高仿真可信度。第三是红外合成数据生成，为目标识别、跟踪、融合感知算法提供可控、可标注、可批量生成的数据。第四是自动评测，自动生成测试用例、解析日志、发现异常并输出验收报告。",
    )
    rows = [
        ("场景自动生成", "根据文字任务、模板和历史场景快速生成仿真工况。", "减少售前和试验准备时间，适合低空和航天批量场景。"),
        ("参数反演校准", "结合实测图像反推温度、材质、大气和传感器参数。", "提升红外仿真可信度，为军工验收提供依据。"),
        ("合成数据服务", "生成红外目标识别、跟踪、融合感知训练数据。", "服务算法团队，也可作为对外增值服务。"),
        ("自动测试与报告", "自动生成测试用例、对比输出、解析日志、生成报告。", "降低交付成本，提高产品可验收性。"),
    ]
    add_matrix(doc, ["AI 能力", "说明", "产品价值"], rows, [1800, 4100, 3460], header_fill=LIGHT_BLUE)

    add_heading(doc, "十二、风险与应对", 1)
    rows = [
        ("性能风险", "ARM 端 GPU 性能不足，复杂场景实时帧率不达标。", "分层渲染、LOD、GPU 并行、分辨率策略、FPGA 协同接口预留。"),
        ("精度风险", "红外仿真效果缺少实测数据校准，客户难以认可。", "建立材料/目标/大气参数库，联合试验单位形成仿真-实测对照数据。"),
        ("需求风险", "客户把平台产品变成定制项目，导致复用性下降。", "区分平台标准能力和项目定制能力，建立版本边界和报价边界。"),
        ("集成风险", "AFSIM、VR、UE5、雷达成像等协同节奏不一致。", "先松耦合接口，再逐步统一坐标、时间、事件和场景资产。"),
        ("合规风险", "开源协议、军工保密、知识产权和国产化适配要求复杂。", "建立开源合规清单、交付环境清单、软著/专利/技术秘密布局。"),
    ]
    add_matrix(doc, ["风险", "表现", "应对措施"], rows, [1600, 3800, 3960], header_fill=FILL, font_size=8.7)

    add_heading(doc, "十三、近期行动建议", 1)
    add_paragraph(
        doc,
        "近期建议不急于把所有方向同时展开，而是先完成四件事。第一，明确 V1.0 必做和不做清单，避免把 UE5、雷达、AI 和全链路辐亮度能力全部压进首版。第二，锁定 2-3 家试点客户，优先选择军工或低空场景中需求清晰、能提供验收反馈的客户。第三，建立跨部门资源池，把 AFSIM、VR、UE5、雷达成像、硬件平台和售前资源纳入路线图。第四，启动模型库、校准数据和 AI 数据生成预研，为后续产品壁垒做准备。",
    )
    add_bullets(
        doc,
        [
            "V1.0 边界：红外内核、GUI 基础配置、SDK/API、ARM 适配、基础模型库、日志和报告。",
            "试点场景：军工半实物测试优先，其次选择低空无人机红外感知或商业航天载荷验证场景。",
            "协同机制：建立仿真软件部牵头、售前/硬件/AFSIM/VR/UE5/雷达团队参与的月度评审机制。",
            "知识产权：围绕国产 ARM 红外仿真适配、热点/亮点建模、协议化目标联动、自动化评估布局软著和专利。",
        ],
    )
    add_callout(
        doc,
        "最终建议",
        "以“国产化三维红外成像仿真产品平台”为主线立项，以军工、低空经济、商业航天解决方案包为首批商业化方向；短期做可交付闭环，中期做行业复制，长期做多模态传感器仿真平台。",
        fill=GREEN_FILL,
    )

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
